#!/usr/bin/env python3
"""Build a clean, searchable CEA writing corpus without changing source files."""

from __future__ import annotations

import argparse
import hashlib
import json
import re
from collections import Counter, defaultdict
from difflib import SequenceMatcher
from pathlib import Path


BLOCK_TAGS = {
    "p", "div", "section", "li", "h1", "h2", "h3", "h4", "blockquote",
    "table", "tr", "br", "figure", "figcaption",
}
MOJIBAKE_MARKERS = ("Ã", "Â", "Ê", "Á", "Ë", "Ô", "‰", "„", "å", "æ", "ç", "ä", "ï¼", "ã€")


def text_score(value: str) -> int:
    cjk = sum("\u4e00" <= char <= "\u9fff" for char in value)
    bad = sum(value.count(marker) for marker in MOJIBAKE_MARKERS)
    replacement = value.count("�")
    return cjk * 3 - bad * 5 - replacement * 30


def decode_chunks(value: str, encoding: str) -> str:
    result: list[str] = []
    buffer: list[str] = []

    def flush() -> None:
        if not buffer:
            return
        original = "".join(buffer)
        buffer.clear()
        try:
            # Some archived pages already contain replacement characters in the
            # middle of an otherwise recoverable UTF-8 byte sequence. Decode
            # the surrounding sequence and retain an explicit replacement mark
            # instead of discarding the entire readable chunk.
            candidate = original.encode(encoding).decode("utf-8", errors="replace")
        except UnicodeEncodeError:
            result.append(original)
            return
        result.append(candidate if text_score(candidate) > text_score(original) else original)

    for char in value:
        try:
            char.encode(encoding)
            buffer.append(char)
        except UnicodeEncodeError:
            flush()
            result.append(char)
    flush()
    return "".join(result)


def repair_mojibake(value: str) -> tuple[str, list[str]]:
    applied: list[str] = []

    def repair_piece(piece: str) -> str:
        repaired = piece
        for _ in range(2):
            candidates = [(text_score(repaired), repaired, None)]
            for encoding in ("mac_roman", "latin1", "cp1252"):
                candidate = decode_chunks(repaired, encoding)
                candidates.append((text_score(candidate), candidate, encoding))
            _, best, encoding = max(candidates, key=lambda item: item[0])
            if best == repaired:
                break
            repaired = best
            if encoding:
                applied.append(encoding)
        return repaired

    # Pages can mix correctly decoded blocks with two different mojibake paths.
    # Repair line-by-line so one damaged paragraph does not prevent recovery of
    # otherwise valid paragraphs elsewhere in the same article.
    repaired = "\n".join(repair_piece(line) for line in value.split("\n"))
    return repaired, sorted(set(applied))


def slugify(value: str, fallback: str) -> str:
    cleaned = re.sub(r"^\d+[_.、。\s-]*", "", value)
    cleaned = re.sub(r"【[^】]+】", "", cleaned)
    cleaned = re.sub(r"[^0-9A-Za-z\u4e00-\u9fff]+", "-", cleaned).strip("-")
    if not cleaned:
        cleaned = fallback
    digest = hashlib.sha1(value.encode("utf-8")).hexdigest()[:8]
    return f"{cleaned[:48]}-{digest}"


def normalize_title(value: str) -> str:
    value = re.sub(r"^\d+[_.、。\s-]*", "", value)
    value = re.sub(r"【[^】]+】", "", value)
    return re.sub(r"[^0-9a-z\u4e00-\u9fff]", "", value.lower())


def node_stream(element) -> tuple[str, list[dict]]:
    chunks: list[str] = []
    images: list[dict] = []

    def walk(node) -> None:
        if node.text:
            chunks.append(node.text)
        for child in node:
            tag = child.tag.lower() if isinstance(child.tag, str) else ""
            if tag == "img":
                source = child.get("data-src") or child.get("src") or ""
                alt = child.get("alt") or "历史文章图片"
                images.append({"src": source, "alt": alt, "data_type": child.get("data-type")})
                chunks.append(f"\n![{alt}]({source})\n")
            else:
                walk(child)
            if child.tail:
                chunks.append(child.tail)
            if tag in BLOCK_TAGS:
                chunks.append("\n")

    walk(element)
    return "".join(chunks), images


def normalize_lines(value: str) -> str:
    value = value.replace("\r", "\n").replace("\u00a0", " ")
    lines = []
    for raw in value.splitlines():
        line = re.sub(r"[ \t]+", " ", raw).strip()
        if not line:
            if lines and lines[-1] != "":
                lines.append("")
            continue
        if not lines or line != lines[-1]:
            lines.append(line)
    return "\n".join(lines).strip()


def strip_repeated_chrome(value: str, category: str) -> tuple[str, str, list[str]]:
    notes: list[str] = []
    footer = ""
    footer_patterns = (
        "全欧/全英中国经济学会（Chinese Economic Association",
        "全英/全欧中国经济学会（Chinese Economic Association",
        "CEA学会出版英文期刊Journal of Chinese Economic and Business Studies",
    )
    positions = [value.rfind(pattern) for pattern in footer_patterns]
    footer_candidates = [position for position in positions if position > len(value) * 0.45]
    footer_at = min(footer_candidates) if footer_candidates else -1
    if footer_at >= 0:
        footer = value[footer_at:].strip()
        value = value[:footer_at].rstrip()
        notes.append("separated_repeated_cea_footer")

    if category == "文章解读":
        nav_at = value.find("往期精彩文章：")
        article_at = value.find("文章题目：")
        if 0 <= nav_at < 1200 and article_at > nav_at:
            value = value[article_at:]
            notes.append("removed_previous_articles_navigation")
    return value, footer, notes


def extract_html(path: Path, html_root: Path, output: Path, style_profile) -> dict:
    try:
        from lxml import html as lxml_html
    except ImportError as exc:
        raise SystemExit("lxml is required; use the Codex bundled Python runtime") from exc

    document = lxml_html.fromstring(path.read_bytes())
    matches = document.xpath('//*[@id="js_content"]')
    category = path.parent.name
    title = re.sub(r"^\d+[_.、。\s-]*", "", path.stem)
    if not matches:
        return {
            "kind": "html", "category": category, "title": title,
            "source_path": str(path.resolve()), "status": "missing_js_content",
        }

    body = matches[0]
    raw_text, images = node_stream(body)
    repaired, encodings = repair_mojibake(raw_text)
    repaired = normalize_lines(repaired)
    repaired, footer, chrome_notes = strip_repeated_chrome(repaired, category)
    remaining_bad = repaired.count("�") + sum(repaired.count(marker) for marker in MOJIBAKE_MARKERS)
    bad_ratio = remaining_bad / max(1, len(repaired))
    text_quality = "degraded_mojibake" if repaired.count("�") >= 5 or bad_ratio >= 0.01 else "usable"

    for image in images:
        image["alt"], _ = repair_mojibake(image["alt"])

    for node in body.iterdescendants():
        style = node.get("style") or ""
        for color in re.findall(r"(?:color|background(?:-color)?)\s*:\s*(#[0-9a-fA-F]{3,8}|rgba?\([^)]*\))", style):
            style_profile[category]["colors"][color.lower()] += 1
        for size in re.findall(r"font-size\s*:\s*([^;]+)", style):
            style_profile[category]["font_sizes"][size.strip()] += 1
        for alignment in re.findall(r"text-align\s*:\s*([^;]+)", style):
            style_profile[category]["alignments"][alignment.strip()] += 1

    slug = slugify(title, "article")
    relative_file = Path("html") / category / f"{slug}.md"
    target = output / relative_file
    target.parent.mkdir(parents=True, exist_ok=True)
    header = (
        f"<!-- source: {path.resolve()} -->\n"
        f"<!-- category: {category} -->\n"
        f"<!-- historical-example-only: true -->\n\n"
        f"# {title}\n\n"
    )
    target.write_text(header + repaired + "\n", encoding="utf-8")
    return {
        "kind": "html", "category": category, "title": title,
        "source_path": str(path.resolve()), "file": str(relative_file),
        "status": "ok", "characters": len(repaired), "images": images,
        "footer_detected": bool(footer), "footer_characters": len(footer),
        "footer_excerpt": footer[:2400],
        "encoding_repairs": encodings, "cleanup": chrome_notes,
        "text_quality": text_quality, "remaining_mojibake_ratio": round(bad_ratio, 6),
    }


def extract_word(path: Path, output: Path) -> dict:
    try:
        from docx import Document
    except ImportError as exc:
        raise SystemExit("python-docx is required; use the Codex bundled Python runtime") from exc

    document = Document(path)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    tables = []
    for table in document.tables:
        rows = []
        for row in table.rows:
            rows.append(" | ".join(cell.text.strip().replace("\n", " / ") for cell in row.cells))
        if rows:
            tables.append("\n".join(rows))
    text = "\n\n".join(paragraphs)
    if tables:
        text += "\n\n## 表格抽取\n\n" + "\n\n".join(tables)

    # The folder name carries the editorial title more reliably than short names
    # such as BERT.docx or 会议预告.docx.
    title = path.parent.name.strip()
    category_match = re.search(r"【([^】]+)】", path.parent.name)
    category = category_match.group(1) if category_match else ("会议通知" if "会议" in path.parent.name else "未分类")
    if "文章解读+技术分享" in category:
        category = "学术前沿"
    slug = slugify(path.parent.name + title, "word")
    relative_file = Path("word") / f"{slug}.md"
    target = output / relative_file
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text(
        f"<!-- source: {path.resolve()} -->\n<!-- category: {category} -->\n"
        f"<!-- historical-example-only: true -->\n\n# {title}\n\n{text}\n",
        encoding="utf-8",
    )
    image_count = sum(1 for rel in document.part.rels.values() if "image" in rel.reltype)
    related_pdfs = [str(item.resolve()) for item in path.parent.glob("*.pdf")]
    return {
        "kind": "word", "category": category, "title": title,
        "source_path": str(path.resolve()), "file": str(relative_file),
        "status": "ok", "characters": len(text), "embedded_images": image_count,
        "tables": len(document.tables), "related_pdfs": related_pdfs,
    }


def content_grams(text: str, width: int = 10) -> set[str]:
    normalized = re.sub(r"[^0-9a-z\u4e00-\u9fff]", "", text.lower())
    return {normalized[index:index + width] for index in range(0, max(0, len(normalized) - width + 1), width)}


def build_pairs(entries: list[dict], output: Path) -> list[dict]:
    html_entries = [entry for entry in entries if entry.get("kind") == "html" and entry.get("status") == "ok"]
    word_entries = [entry for entry in entries if entry.get("kind") == "word" and entry.get("status") == "ok"]
    pairs = []
    for word in word_entries:
        word_text = (output / word["file"]).read_text(encoding="utf-8")
        grams = content_grams(word_text)
        candidates = []
        for html in html_entries:
            html_text = (output / html["file"]).read_text(encoding="utf-8")
            coverage = sum(gram in html_text for gram in grams) / max(1, len(grams))
            title_score = SequenceMatcher(None, normalize_title(word["title"]), normalize_title(html["title"])).ratio()
            score = 0.75 * coverage + 0.25 * title_score
            candidates.append((score, coverage, title_score, html))
        if not candidates:
            continue
        score, coverage, title_score, html = max(candidates, key=lambda item: item[0])
        if score >= 0.16:
            pairs.append({
                "word_file": word["file"], "html_file": html["file"],
                "word_title": word["title"], "html_title": html["title"],
                "score": round(score, 4), "content_coverage": round(coverage, 4),
                "title_similarity": round(title_score, 4), "status": "candidate_needs_review",
            })
    return pairs


def serialize_style_profile(profile) -> dict:
    result = {}
    for category, groups in profile.items():
        result[category] = {
            name: counter.most_common(20) for name, counter in groups.items()
        }
    return result


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--word-dir", type=Path, required=True)
    parser.add_argument("--html-dir", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()

    if not args.word_dir.is_dir() or not args.html_dir.is_dir():
        parser.error("--word-dir and --html-dir must be existing directories")
    output = args.output.resolve()
    output.mkdir(parents=True, exist_ok=True)

    profile = defaultdict(lambda: {
        "colors": Counter(), "font_sizes": Counter(), "alignments": Counter(),
    })
    entries: list[dict] = []
    for path in sorted(args.html_dir.rglob("*.html")):
        entries.append(extract_html(path, args.html_dir, output, profile))
    for path in sorted(args.word_dir.rglob("*.docx")):
        if not path.name.startswith("._"):
            entries.append(extract_word(path, output))

    # Remove only stale Markdown files previously generated inside the two
    # bounded corpus subdirectories. Never touch source material or other
    # files under references/corpus.
    current_files = {(output / entry["file"]).resolve() for entry in entries if entry.get("file")}
    for subdirectory in (output / "html", output / "word"):
        if subdirectory.exists():
            for generated in subdirectory.rglob("*.md"):
                if generated.resolve() not in current_files:
                    generated.unlink()

    pairs = build_pairs(entries, output)
    index = {
        "schema_version": 1,
        "notice": "Historical examples control style and structure only; verify all current facts from original or official sources.",
        "entries": entries,
    }
    (output / "index.json").write_text(json.dumps(index, ensure_ascii=False, indent=2), encoding="utf-8")
    (output / "pairs.json").write_text(json.dumps(pairs, ensure_ascii=False, indent=2), encoding="utf-8")
    (output / "style-profile.json").write_text(
        json.dumps(serialize_style_profile(profile), ensure_ascii=False, indent=2), encoding="utf-8"
    )
    html_entries = [entry for entry in entries if entry["kind"] == "html"]
    word_entries = [entry for entry in entries if entry["kind"] == "word"]
    summary = {
        "html_total": sum(entry["kind"] == "html" for entry in entries),
        "html_clean": sum(entry["kind"] == "html" and entry.get("status") == "ok" for entry in entries),
        "word_total": sum(entry["kind"] == "word" for entry in entries),
        "candidate_pairs": len(pairs),
        "html_text_quality": dict(Counter(entry.get("text_quality", "unavailable") for entry in html_entries)),
        "html_categories": dict(Counter(entry.get("category", "unknown") for entry in html_entries)),
        "html_with_cea_footer": sum(bool(entry.get("footer_detected")) for entry in html_entries),
        "word_categories": dict(Counter(entry.get("category", "unknown") for entry in word_entries)),
    }
    (output / "summary.json").write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(summary, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
